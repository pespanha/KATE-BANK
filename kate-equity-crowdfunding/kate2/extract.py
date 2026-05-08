import docx
import fitz # PyMuPDF
import sys

def extract_docx(filepath, out_filepath):
    doc = docx.Document(filepath)
    with open(out_filepath, 'w', encoding='utf-8') as f:
        for p in doc.paragraphs:
            f.write(p.text + '\n')
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    f.write(cell.text + '\t')
                f.write('\n')

def extract_pdf(filepath, out_filepath):
    doc = fitz.open(filepath)
    with open(out_filepath, 'w', encoding='utf-8') as f:
        for page in doc:
            f.write(page.get_text())

print("Extracting docx...")
extract_docx("Stellar37 - ECF Stellar (1).docx", "docx_text.txt")
print("Extracting pdf...")
extract_pdf("MANUAL DA MARCA_Kate.pdf", "pdf_text.txt")
print("Done.")
